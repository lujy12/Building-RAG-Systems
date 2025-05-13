import os
import fitz
import docx
import faiss
import pickle
import numpy as np
from pathlib import Path
from typing import List, Dict
from sklearn.metrics.pairwise import cosine_similarity
from langchain.text_splitter import CharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document
from transformers import pipeline


def load_documents(directory: str) -> List[Dict]:
    documents = []
    for file_path in Path(directory).glob("*"):
        try:
            if file_path.suffix.lower() == ".pdf":
                text = ""
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text()
            elif file_path.suffix.lower() == ".docx":
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif file_path.suffix.lower() == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                print(f"Unsupported file type: {file_path.name}")
                continue

            documents.append({
                "text": text,
                "metadata": {
                    "filename": file_path.name,
                    "path": str(file_path.resolve())
                }
            })

        except Exception as e:
            print(f"Error reading {file_path.name}: {e}")
    return documents


def split_documents(docs: List[Dict], chunk_size=500, chunk_overlap=50) -> List[Dict]:
    splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )

    chunks = []
    for doc in docs:
        text_chunks = splitter.split_text(doc["text"])
        for i, chunk in enumerate(text_chunks):
            chunks.append({
                "text": chunk,
                "metadata": {
                    **doc["metadata"],
                    "chunk": i + 1,
                    "text": chunk
                }
            })
    return chunks



def embed_and_store(chunks: List[Dict], model_name: str, faiss_index_path: str):
    """Supports embedding with multiple models. e.g., MiniLM and Paraphrase-MiniLM"""
    
    print(f"\nLoading embedding model: {model_name}")
    model = SentenceTransformer(model_name)

    texts = [chunk["text"] for chunk in chunks]
    metadatas = [chunk["metadata"] for chunk in chunks]

    print("Generating embeddings...")
    embeddings = model.encode(texts, show_progress_bar=True)

    print("Building FAISS index...")
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(embeddings))

    faiss.write_index(index, faiss_index_path + ".index")
    with open(faiss_index_path + "_metadata.pkl", "wb") as f:
        pickle.dump(metadatas, f)

    print(f"FAISS index saved to: {faiss_index_path}.index")
    print(f"Metadata saved to: {faiss_index_path}_metadata.pkl")


def load_faiss_index(index_path: str):
    print("Loading FAISS index and metadata...")
    index = faiss.read_index(index_path + ".index")
    with open(index_path + "_metadata.pkl", "rb") as f:
        metadata = pickle.load(f)
    return index, metadata

def search_documents(query: str, model_name: str, index_path: str, k=3):
    model = SentenceTransformer(model_name)
    query_embedding = model.encode([query])
    index, metadata = load_faiss_index(index_path)
    distances, indices = index.search(np.array(query_embedding), k)

    print(f"\nTop {k} Results for Query: \"{query}\"")
    for i, idx in enumerate(indices[0]):
        print(f"\nResult #{i+1} (Distance: {distances[0][i]:.4f})")
        print(f"File: {metadata[idx]['filename']} | Chunk: {metadata[idx]['chunk']}")
        print("Content Preview:")
        print("-" * 40)
        print(metadata[idx]['text'][:500] + "...\n")



def mmr(query_embedding, doc_embeddings, k=3, lambda_param=0.5):
    selected = []
    unselected = list(range(len(doc_embeddings)))
    query_embedding = np.array(query_embedding).reshape(1, -1)
    doc_embeddings = np.array(doc_embeddings)

    sim_to_query = cosine_similarity(query_embedding, doc_embeddings)[0]
    sim_between_docs = cosine_similarity(doc_embeddings)

    for _ in range(k):
        mmr_scores = []
        for idx in unselected:
            diversity = 0 if not selected else max(sim_between_docs[idx][j] for j in selected)
            mmr_score = lambda_param * sim_to_query[idx] - (1 - lambda_param) * diversity
            mmr_scores.append((mmr_score, idx))

        mmr_scores.sort(reverse=True)
        best = mmr_scores[0][1]
        selected.append(best)
        unselected.remove(best)

    return selected

def search_documents_mmr(query: str, model_name: str, index_path: str, k=3):
    model = SentenceTransformer(model_name)
    query_embedding = model.encode([query])
    index, metadata = load_faiss_index(index_path)
    all_embeddings = np.array([index.reconstruct(i) for i in range(index.ntotal)])
    selected_indices = mmr(query_embedding, all_embeddings, k=k)

    print(f"\nMMR Top {k} Diverse Results for Query: \"{query}\"")
    for rank, idx in enumerate(selected_indices):
        print(f"\nResult #{rank+1}")
        print(f"File: {metadata[idx]['filename']} | Chunk: {metadata[idx]['chunk']}")
        print("Content Preview:")
        print("-" * 40)
        print(metadata[idx]['text'][:500] + "...\n")



def generate_answer_with_huggingface(query: str, index_path: str, model_name="google/flan-t5-base"):
    print("\nLoading FAISS retriever and generating answer (HuggingFace)...")
    index = faiss.read_index(index_path + ".index")
    with open(index_path + "_metadata.pkl", "rb") as f:
        metadata = pickle.load(f)

    docs = [Document(page_content=meta["text"], metadata=meta) for meta in metadata]
    embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    db = FAISS.from_documents(docs, embedding_model)
    retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 3})

    retrieved_docs = retriever.get_relevant_documents(query)
    context = "\n".join([doc.page_content for doc in retrieved_docs])

    prompt = f"""You are a helpful assistant. Use the context below to answer the question.

Context:
{context}

Question:
{query}

Answer:"""

    generator = pipeline("text2text-generation", model=model_name, tokenizer=model_name)
    response = generator(prompt, max_new_tokens=300)[0]['generated_text']

    print("\nAnswer:")
    print("-" * 50)
    print(response.strip())


def evaluate_retrieval_system(index_path: str, model_name: str, queries: List[Dict], k=3):
    from sklearn.metrics import precision_score, recall_score, f1_score

    print(f"\nEvaluating retrieval for model: {model_name}")
    model = SentenceTransformer(model_name)
    index, metadata = load_faiss_index(index_path)
    all_embeddings = np.array([index.reconstruct(i) for i in range(index.ntotal)])

    y_true_all = []
    y_pred_all = []

    for item in queries:
        query = item['query']
        expected_keywords = item['keywords'] 

        query_embedding = model.encode([query])
        distances, indices = index.search(np.array(query_embedding), k)
        retrieved_texts = [metadata[idx]['text'] for idx in indices[0]]

        true = []
        pred = []
        for keyword in expected_keywords:
            true.append(1)
            if any(keyword.lower() in chunk.lower() for chunk in retrieved_texts):
                pred.append(1)
            else:
                pred.append(0)

        y_true_all.extend(true)
        y_pred_all.extend(pred)

    precision = precision_score(y_true_all, y_pred_all)
    recall = recall_score(y_true_all, y_pred_all)
    f1 = f1_score(y_true_all, y_pred_all)

    print(f"Precision: {precision:.2f}, Recall: {recall:.2f}, F1-Score: {f1:.2f}")
    return precision, recall, f1


def compare_models(chunks: List[Dict], queries: List[Dict]):
    os.makedirs("faiss_store", exist_ok=True)
    configs = {
        "MiniLM": {
            "model_name": "all-MiniLM-L6-v2",
            "index_path": "faiss_store/index_miniLM"
        },
        "Paraphrase-MiniLM": {
            "model_name": "paraphrase-MiniLM-L6-v2",
            "index_path": "faiss_store/index_paraLM"
        }
    }

    for label, cfg in configs.items():
        print(f"\nEmbedding with {label}...")
        embed_and_store(chunks, model_name=cfg["model_name"], faiss_index_path=cfg["index_path"])

    print("\n=== Evaluation ===")
    for label, cfg in configs.items():
        evaluate_retrieval_system(
            index_path=cfg["index_path"],
            model_name=cfg["model_name"],
            queries=queries
        )


if __name__ == "__main__":
    mode = input("Select mode: [1] Build Index  [2] Search  [3] MMR Search  [4] RAG Answer  [5] Compare Models: ").strip()

    if mode == "1":
        docs = load_documents("documents")
        print(f"\nLoaded {len(docs)} document(s).")
        chunks = split_documents(docs, chunk_size=500, chunk_overlap=50)
        print(f"Split into {len(chunks)} chunks.")
        os.makedirs("faiss_store", exist_ok=True)
        embed_and_store(
            chunks=chunks,
            model_name="all-MiniLM-L6-v2",
            faiss_index_path="faiss_store/index_miniLM"
        )

    elif mode == "2":
        query = input("\nEnter your query: ")
        search_documents(
            query=query,
            model_name="all-MiniLM-L6-v2",
            index_path="faiss_store/index_miniLM",
            k=3
        )

    elif mode == "3":
        query = input("\nEnter your query for MMR search: ")
        search_documents_mmr(
            query=query,
            model_name="all-MiniLM-L6-v2",
            index_path="faiss_store/index_miniLM",
            k=3
        )

    elif mode == "4":
        query = input("\nEnter your question for RAG LLM answer: ")
        generate_answer_with_huggingface(
            query=query,
            index_path="faiss_store/index_miniLM",
            model_name="google/flan-t5-base"
        )

    
    elif mode == "5":
        # Compare models
        docs = load_documents("documents")
        chunks = split_documents(docs, chunk_size=500, chunk_overlap=50)
        test_queries = [
            {"query": "What is LangChain used for?", "keywords": ["LangChain"]},
            {"query": "Explain semantic similarity search", "keywords": ["semantic", "similarity"]},
            {"query": "What is FAISS used for?", "keywords": ["FAISS"]}
        ]
        compare_models(chunks, queries=test_queries)

    else:
        print("Invalid option. Choose 1, 2, 3, or 4.")

